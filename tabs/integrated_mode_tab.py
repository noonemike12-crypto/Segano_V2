import os
import base64
import uuid
import ffmpeg
import gnupg
from docx import Document
from datetime import datetime
from PyQt5.QtWidgets import (
    QWidget, QVBoxLayout, QHBoxLayout, QGroupBox, QLabel, 
    QPushButton, QComboBox, QTextEdit, QFrame, QFileDialog, 
    QTableWidget, QTableWidgetItem, QHeaderView, QProgressBar, QTabWidget, QPlainTextEdit,
    QMessageBox
)
from PyQt5.QtCore import Qt

from utils.encryption import CryptoUtils
from utils.steganography import hide_lsb_image, extract_lsb_image

class IntegrationTab(QWidget):
    def __init__(self):
        super().__init__()
        self.selected_files = []
        self.init_ui()

    def init_ui(self):
        main_layout = QVBoxLayout(self)
        main_layout.setSpacing(20)
        main_layout.setContentsMargins(30, 30, 30, 30)

        # --- Top Section: Intro ---
        intro_layout = QVBoxLayout()
        title = QLabel("Integrated Steganography Engine")
        title.setStyleSheet("font-size: 18pt; font-weight: 800; color: white;")
        subtitle = QLabel("Combine multiple encryption and steganography techniques for maximum security and obfuscation.")
        subtitle.setObjectName("subTitle")
        intro_layout.addWidget(title)
        intro_layout.addWidget(subtitle)
        main_layout.addLayout(intro_layout)

        # --- Middle Section: Main Grid ---
        grid_layout = QHBoxLayout()
        grid_layout.setSpacing(30)

        # Left Column: Configuration & Files
        left_col = QVBoxLayout()
        
        # Mode Selection
        mode_group = QGroupBox("STRATEGY SELECTION")
        mode_layout = QVBoxLayout(mode_group)
        self.mode_selector = QComboBox()
        self.modes = [
            "🔄 Mode 1: AES + Split Payload (Image + Audio)",
            "📄 Mode 2: DOCX + RSA + Video Metadata",
            "🎛️ Mode 3: AES + Triple Split (Image + Audio + Video)",
            "🧬 Mode 4: AES + RSA + Metadata Injection",
            "🧫 Mode 5: GPG + Metadata + EOF Padding",
            "🧩 Mode 6: AES + LSB + Metadata + Checksum",
            "🔄 Mode 7: Multi-Layer Transform + Obfuscation",
            "🧠 Mode 8: AES + GPG + Cross-Media Sync",
            "🌀 Mode 9: Nested Stego (Recursive Hiding)",
            "🧾 Mode 10: Split + Layered + Temporal Lock"
        ]
        self.mode_selector.addItems(self.modes)
        mode_layout.addWidget(self.mode_selector)
        
        self.mode_desc = QLabel("Select a strategy to view technical details...")
        self.mode_desc.setWordWrap(True)
        self.mode_desc.setStyleSheet("color: #94a3b8; font-size: 9pt; font-style: italic;")
        mode_layout.addWidget(self.mode_desc)
        left_col.addWidget(mode_group)

        # File Management
        file_group = QGroupBox("CARRIER MANAGEMENT")
        file_layout = QVBoxLayout(file_group)
        
        file_btns = QHBoxLayout()
        self.add_file_btn = QPushButton("➕ Add Carriers")
        self.add_file_btn.clicked.connect(self.add_files)
        self.clear_files_btn = QPushButton("🗑️ Clear All")
        self.clear_files_btn.setObjectName("dangerBtn")
        self.clear_files_btn.clicked.connect(self.clear_files)
        file_btns.addWidget(self.add_file_btn)
        file_btns.addWidget(self.clear_files_btn)
        file_layout.addLayout(file_btns)
        
        self.files_table = QTableWidget()
        self.files_table.setColumnCount(3)
        self.files_table.setHorizontalHeaderLabels(["Filename", "Type", "Size"])
        self.files_table.horizontalHeader().setSectionResizeMode(QHeaderView.Stretch)
        self.files_table.setMinimumHeight(150)
        file_layout.addWidget(self.files_table)
        left_col.addWidget(file_group)
        
        grid_layout.addLayout(left_col, 3)

        # Right Column: Payload & Execution
        right_col = QVBoxLayout()
        
        # Payload Section
        payload_group = QGroupBox("PAYLOAD INPUT")
        payload_layout = QVBoxLayout(payload_group)
        self.text_input = QPlainTextEdit()
        self.text_input.setPlaceholderText("Enter the sensitive data to be processed...")
        self.text_input.setMinimumHeight(150)
        payload_layout.addWidget(self.text_input)
        right_col.addWidget(payload_group)

        # Execution Control
        exec_group = QGroupBox("ENGINE CONTROL")
        exec_layout = QVBoxLayout(exec_group)
        self.execute_btn = QPushButton("🚀 EXECUTE STRATEGY")
        self.execute_btn.setObjectName("primaryBtn")
        self.execute_btn.setMinimumHeight(50)
        self.execute_btn.clicked.connect(self.process_integrated)
        exec_layout.addWidget(self.execute_btn)
        right_col.addWidget(exec_group)
        
        grid_layout.addLayout(right_col, 2)
        main_layout.addLayout(grid_layout)

        # Status & Progress
        status_group = QGroupBox("SYSTEM STATUS")
        status_layout = QVBoxLayout(status_group)
        self.progress_bar = QProgressBar()
        self.progress_bar.setValue(0)
        status_layout.addWidget(self.progress_bar)
        
        self.log_output = QTextEdit()
        self.log_output.setReadOnly(True)
        self.log_output.setMaximumHeight(100)
        self.log_output.setPlaceholderText("Engine logs...")
        status_layout.addWidget(self.log_output)
        main_layout.addWidget(status_group)

        self.mode_selector.currentIndexChanged.connect(self.update_desc)
        self.update_desc()

    def update_desc(self):
        descs = [
            "🔐 เข้ารหัส AES → แบ่งครึ่ง → ซ่อนในภาพและเสียง (LSB)",
            "📄 สร้าง DOCX → RSA → ซ่อนใน Metadata วิดีโอ",
            "🎛️ AES → แบ่ง 3 ส่วน → ซ่อนในภาพ, เสียง และวิดีโอ",
            "🧬 AES + RSA → ซ่อนใน Metadata หลายไฟล์",
            "🧫 GPG Encryption → ซ่อนใน Metadata + EOF",
            "🧩 AES + LSB + Metadata + Checksum (ความปลอดภัยสูง)",
            "🔄 แปลงหลายชั้น (Base64+Gzip) → ซ่อนหลายจุด",
            "🧠 กระจายข้อมูลและกุญแจ (AES+GPG) ในหลายไฟล์",
            "🌀 ซ่อนซ้อนกันเป็นชั้นๆ (Nested Steganography)",
            "🧾 แบ่งส่วนข้อมูล + เข้ารหัสหลายชั้น + ระบบล็อกเวลา"
        ]
        self.mode_desc.setText(descs[self.mode_selector.currentIndex()])

    def add_files(self):
        paths, _ = QFileDialog.getOpenFileNames(self, "เลือกไฟล์")
        if paths:
            for p in paths:
                self.selected_files.append(p)
                row = self.files_table.rowCount()
                self.files_table.insertRow(row)
                self.files_table.setItem(row, 0, QTableWidgetItem(os.path.basename(p)))
                self.files_table.setItem(row, 1, QTableWidgetItem(os.path.splitext(p)[1]))
                self.files_table.setItem(row, 2, QTableWidgetItem(f"{os.path.getsize(p)/1024:.1f} KB"))

    def clear_files(self):
        self.selected_files = []
        self.files_table.setRowCount(0)

    def process_integrated(self):
        mode = self.mode_selector.currentIndex()
        text = self.text_input.toPlainText()
        if not text or not self.selected_files:
            QMessageBox.warning(self, "คำเตือน", "กรุณากรอกข้อความและเลือกไฟล์ที่จำเป็น")
            return
        
        self.log_output.append(f"🔄 กำลังเริ่ม {self.modes[mode]}...")
        
        try:
            if mode == 0: # โหมด 1: AES + แบ่งครึ่ง (ภาพ + เสียง)
                if len(self.selected_files) < 2:
                    raise ValueError("โหมดนี้ต้องการอย่างน้อย 2 ไฟล์ (ภาพ และ เสียง)")
                
                self.log_output.append("🔐 เข้ารหัส AES...")
                key = "sieng_secret_key_32_chars_long!!!" 
                encrypted = CryptoUtils.aes_encrypt(text, key)
                
                half = len(encrypted) // 2
                p1, p2 = encrypted[:half], encrypted[half:]
                
                self.log_output.append("🖼️ ซ่อนส่วนที่ 1 ในภาพ...")
                out_img = os.path.join("photoexample", "output", "mode1_part1.png")
                os.makedirs(os.path.dirname(out_img), exist_ok=True)
                hide_lsb_image(self.selected_files[0], p1, out_img)
                
                self.log_output.append("🎵 ซ่อนส่วนที่ 2 ในเสียง...")
                # ใช้ logic ซ่อนในไฟล์เสียงแบบง่าย (EOF Append สำหรับโหมดรวมเพื่อความรวดเร็วและรองรับหลายฟอร์แมต)
                out_audio = os.path.join("audioexample", "output", "mode1_part2" + os.path.splitext(self.selected_files[1])[1])
                os.makedirs(os.path.dirname(out_audio), exist_ok=True)
                with open(self.selected_files[1], 'rb') as f_in, open(out_audio, 'wb') as f_out:
                    f_out.write(f_in.read())
                    f_out.write(b"SIENG_START" + p2.encode('utf-8') + b"SIENG_END")
                
                self.log_output.append(f"✅ สำเร็จ!\n- ส่วนที่ 1: {out_img}\n- ส่วนที่ 2: {out_audio}")

            elif mode == 1: # โหมด 2: DOCX + RSA + Video Metadata
                if not self.selected_files:
                    raise ValueError("โหมดนี้ต้องการไฟล์วิดีโออย่างน้อย 1 ไฟล์")
                
                self.log_output.append("🔑 กำลังสร้างกุญแจ RSA...")
                priv_key, pub_key = CryptoUtils.rsa_generate_keys()
                
                # บันทึกกุญแจไว้ในโฟลเดอร์ผลลัพธ์
                keys_dir = os.path.join("output_files", "keys")
                os.makedirs(keys_dir, exist_ok=True)
                with open(os.path.join(keys_dir, "mode2_private.pem"), "w") as f: f.write(priv_key)
                with open(os.path.join(keys_dir, "mode2_public.pem"), "w") as f: f.write(pub_key)
                
                self.log_output.append("🔐 เข้ารหัสข้อความด้วย RSA...")
                encrypted = CryptoUtils.rsa_encrypt(text, pub_key)
                
                self.log_output.append("📄 สร้างไฟล์ DOCX...")
                doc = Document()
                doc.add_heading('SIENG PRO Secure Document', 0)
                doc.add_paragraph(f"Generated at: {datetime.now()}")
                doc.add_paragraph("This document contains encrypted data.")
                doc.add_paragraph(encrypted)
                
                out_docx = os.path.join("output_files", "mode2_secure.docx")
                os.makedirs(os.path.dirname(out_docx), exist_ok=True)
                doc.save(out_docx)
                
                self.log_output.append("🎬 ซ่อนใน Video Metadata...")
                video_file = self.selected_files[0]
                out_video = os.path.join("vdio", "output", "mode2_meta_" + os.path.basename(video_file))
                os.makedirs(os.path.dirname(out_video), exist_ok=True)
                
                # ใช้ ffmpeg ซ่อน encrypted text ใน comment metadata
                (
                    ffmpeg
                    .input(video_file)
                    .output(out_video, metadata=f"comment={encrypted}", codec="copy")
                    .overwrite_output()
                    .run()
                )
                
                self.log_output.append(f"✅ สำเร็จ!\n- ไฟล์ DOCX: {out_docx}\n- วิดีโอ (Metadata): {out_video}\n- กุญแจ RSA: {keys_dir}")

            elif mode == 2: # โหมด 3: AES + แบ่ง 3 ส่วน (ภาพ + เสียง + วิดีโอ)
                if len(self.selected_files) < 3:
                    raise ValueError("โหมดนี้ต้องการอย่างน้อย 3 ไฟล์ (ภาพ, เสียง, วิดีโอ)")
                
                self.log_output.append("🔐 เข้ารหัส AES...")
                key = "sieng_secret_key_32_chars_long!!!"
                encrypted = CryptoUtils.aes_encrypt(text, key)
                
                third = len(encrypted) // 3
                p1, p2, p3 = encrypted[:third], encrypted[third:2*third], encrypted[2*third:]
                
                self.log_output.append("🖼️ ซ่อนส่วนที่ 1 ในภาพ...")
                out_img = os.path.join("photoexample", "output", "mode3_part1.png")
                hide_lsb_image(self.selected_files[0], p1, out_img)
                
                self.log_output.append("🎵 ซ่อนส่วนที่ 2 ในเสียง...")
                out_audio = os.path.join("audioexample", "output", "mode3_part2" + os.path.splitext(self.selected_files[1])[1])
                with open(self.selected_files[1], 'rb') as f_in, open(out_audio, 'wb') as f_out:
                    f_out.write(f_in.read())
                    f_out.write(b"SIENG_START" + p2.encode('utf-8') + b"SIENG_END")
                
                self.log_output.append("🎬 ซ่อนส่วนที่ 3 ในวิดีโอ (EOF)...")
                out_video = os.path.join("vdio", "output", "mode3_part3" + os.path.splitext(self.selected_files[2])[1])
                with open(self.selected_files[2], 'rb') as f_in, open(out_video, 'wb') as f_out:
                    f_out.write(f_in.read())
                    f_out.write(b"SIENG_START" + p3.encode('utf-8') + b"SIENG_END")
                
                self.log_output.append(f"✅ สำเร็จ!\n- ส่วนที่ 1: {out_img}\n- ส่วนที่ 2: {out_audio}\n- ส่วนที่ 3: {out_video}")

            elif mode == 3: # โหมด 4: AES + RSA + Metadata
                if len(self.selected_files) < 2:
                    raise ValueError("โหมดนี้ต้องการไฟล์อย่างน้อย 2 ไฟล์")
                
                self.log_output.append("🔐 เข้ารหัส AES...")
                aes_key = "aes_key_for_mode4_!!!!"
                encrypted_text = CryptoUtils.aes_encrypt(text, aes_key)
                
                self.log_output.append("🔑 เข้ารหัส AES Key ด้วย RSA...")
                priv_key, pub_key = CryptoUtils.rsa_generate_keys()
                encrypted_key = CryptoUtils.rsa_encrypt(aes_key, pub_key)
                
                self.log_output.append("🎬 ซ่อนข้อมูลใน Metadata หลายไฟล์...")
                out_files = []
                for i, f_path in enumerate(self.selected_files[:2]):
                    out_path = os.path.join("output_files", f"mode4_meta_{i}_{os.path.basename(f_path)}")
                    os.makedirs(os.path.dirname(out_path), exist_ok=True)
                    
                    meta_val = encrypted_text if i == 0 else encrypted_key
                    field = "comment" if i == 0 else "title"
                    
                    try:
                        (
                            ffmpeg
                            .input(f_path)
                            .output(out_path, metadata=f"{field}={meta_val}", codec="copy")
                            .overwrite_output()
                            .run()
                        )
                        out_files.append(out_path)
                    except:
                        with open(f_path, 'rb') as f_in, open(out_path, 'wb') as f_out:
                            f_out.write(f_in.read())
                            f_out.write(f"\nSIENG_META_{field}:{meta_val}".encode())
                        out_files.append(out_path)
                
                self.log_output.append(f"✅ สำเร็จ!\n- ไฟล์ผลลัพธ์: {', '.join(out_files)}")

            elif mode == 4: # โหมด 5: GPG + Metadata + EOF
                if not self.selected_files:
                    raise ValueError("ต้องการไฟล์อย่างน้อย 1 ไฟล์")
                
                self.log_output.append("🔐 เข้ารหัสด้วย GPG (Symmetric)...")
                gpg = gnupg.GPG()
                passphrase = "sieng_gpg_passphrase"
                encrypted = gpg.encrypt(text, recipients=None, symmetric=True, passphrase=passphrase)
                
                if not encrypted.ok:
                    raise ValueError(f"GPG Error: {encrypted.status}")
                
                enc_data = str(encrypted)
                
                self.log_output.append("🎬 ซ่อนใน Metadata และ EOF...")
                f_path = self.selected_files[0]
                out_path = os.path.join("output_files", f"mode5_hybrid_{os.path.basename(f_path)}")
                os.makedirs(os.path.dirname(out_path), exist_ok=True)
                
                half = len(enc_data) // 2
                p1, p2 = enc_data[:half], enc_data[half:]
                
                try:
                    (
                        ffmpeg
                        .input(f_path)
                        .output(out_path, metadata=f"comment={p1}", codec="copy")
                        .overwrite_output()
                        .run()
                    )
                    with open(out_path, 'ab') as f:
                        f.write(b"\nSIENG_GPG_PART2:" + p2.encode())
                except:
                    with open(f_path, 'rb') as f_in, open(out_path, 'wb') as f_out:
                        f_out.write(f_in.read())
                        f_out.write(b"\nSIENG_GPG_FULL:" + enc_data.encode())
                
                self.log_output.append(f"✅ สำเร็จ!\n- ไฟล์ผลลัพธ์: {out_path}\n- Passphrase: {passphrase}")

            else:
                self.log_output.append("⚠️ โหมดอื่นๆ กำลังอยู่ในการพัฒนา")
                
        except Exception as e:
            self.log_output.append(f"❌ ข้อผิดพลาด: {str(e)}")
